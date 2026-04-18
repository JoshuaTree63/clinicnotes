import io
import json
import os

import docx
from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.concurrency import run_in_threadpool
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from services.groq_service import diarize_transcript, format_transcript_for_llm
from services.rag_service import (
    analyze_with_groq,
    analyze_with_openrouter,
    retrieve_context,
)

router = APIRouter()
SESSIONS_DIR = os.getenv("SESSIONS_DIR", "./data/sessions")
AUDIO_DIR = os.getenv("AUDIO_DIR", "./data/audio")

_AUDIO_MIME = {
    ".mp3": "audio/mpeg",
    ".m4a": "audio/mp4",
    ".mp4": "audio/mp4",
    ".wav": "audio/wav",
    ".webm": "audio/webm",
    ".ogg": "audio/ogg",
    ".flac": "audio/flac",
    ".mpga": "audio/mpeg",
    ".mpeg": "audio/mpeg",
}


class AnalyzeRequest(BaseModel):
    session_id: str


class RenameSpeakerRequest(BaseModel):
    old_name: str
    new_name: str


class MergeSpeakerRequest(BaseModel):
    from_name: str
    into_name: str


class RemoveSpeakerRequest(BaseModel):
    name: str


class EditTurnRequest(BaseModel):
    text: str


def _collapse_adjacent_turns(turns: list[dict]) -> list[dict]:
    """Merge consecutive turns with the same speaker into one, preserving start/end timing."""
    out: list[dict] = []
    for t in turns:
        speaker = t.get("speaker")
        text = (t.get("text") or "").strip()
        if not text:
            continue
        start = t.get("start")
        end = t.get("end")
        if out and out[-1].get("speaker") == speaker:
            out[-1]["text"] = (out[-1]["text"] + " " + text).strip()
            if end is not None:
                out[-1]["end"] = end
        else:
            merged = {"speaker": speaker, "text": text}
            if start is not None:
                merged["start"] = start
            if end is not None:
                merged["end"] = end
            out.append(merged)
    return out


def _save_turns(session: dict, turns: list[dict], path: str) -> dict:
    session["transcript_diarized"] = turns
    session["speaker_count"] = len({t.get("speaker") for t in turns})
    with open(path, "w", encoding="utf-8") as f:
        json.dump(session, f, ensure_ascii=False, indent=2)
    return session


@router.post("/analyze")
async def analyze_session(request: AnalyzeRequest):
    session_path = os.path.join(SESSIONS_DIR, f"{request.session_id}.json")
    if not os.path.exists(session_path):
        raise HTTPException(404, detail="Session not found")

    with open(session_path, "r", encoding="utf-8") as f:
        session = json.load(f)

    # Build transcript text for analysis — prefer diarized format, fall back to legacy
    diarized_turns = session.get("transcript_diarized")
    if diarized_turns:
        transcript = format_transcript_for_llm(diarized_turns)
    else:
        # Backwards compat: old sessions with plain "transcript" field
        transcript = session.get("transcript", "")

    if not transcript:
        raise HTTPException(400, detail="Session has no transcript to analyze")

    # Retrieve relevant PDF chunks
    context_chunks = await run_in_threadpool(retrieve_context, transcript)
    if not context_chunks:
        raise HTTPException(
            400, detail="No indexed documents found. Run /api/index first."
        )

    # Try Groq first, fall back to OpenRouter (running in threads)
    try:
        analysis = await run_in_threadpool(analyze_with_groq, transcript, context_chunks)
    except Exception:
        try:
            analysis = await analyze_with_openrouter(transcript, context_chunks)
        except Exception as openrouter_error:
            raise HTTPException(
                503,
                detail=f"Both Groq and OpenRouter failed: {str(openrouter_error)}",
            )

    # Save analysis back to session
    session["analysis"] = analysis
    with open(session_path, "w", encoding="utf-8") as f:
        json.dump(session, f, ensure_ascii=False, indent=2)

    return {"session_id": request.session_id, "analysis": analysis}


@router.get("/sessions")
def list_sessions():
    """List all stored session records."""
    os.makedirs(SESSIONS_DIR, exist_ok=True)
    sessions = []
    for fname in sorted(os.listdir(SESSIONS_DIR), reverse=True):
        if fname.endswith(".json"):
            with open(os.path.join(SESSIONS_DIR, fname), "r") as f:
                s = json.load(f)
                sessions.append(
                    {
                        "id": s["id"],
                        "date": s["date"],
                        "filename": s.get("filename"),
                        "source": s.get("source", "audio"),
                        "speaker_count": s.get("speaker_count"),
                        "has_analysis": s.get("analysis") is not None,
                    }
                )
    return sessions


@router.get("/sessions/{session_id}")
def get_session(session_id: str):
    path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
    if not os.path.exists(path):
        raise HTTPException(404, detail="Session not found")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


@router.get("/sessions/{session_id}/audio")
def stream_session_audio(session_id: str):
    """Stream the persisted audio file for this session (supports Range requests for seeking)."""
    session_path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
    if not os.path.exists(session_path):
        raise HTTPException(404, detail="Session not found")
    with open(session_path, "r", encoding="utf-8") as f:
        session = json.load(f)

    audio_filename = session.get("audio_filename")
    if not audio_filename:
        raise HTTPException(404, detail="No audio stored for this session")

    audio_path = os.path.join(AUDIO_DIR, audio_filename)
    if not os.path.exists(audio_path):
        raise HTTPException(404, detail="Audio file missing on disk")

    ext = os.path.splitext(audio_filename)[1].lower()
    media_type = _AUDIO_MIME.get(ext, "application/octet-stream")
    return FileResponse(audio_path, media_type=media_type, filename=audio_filename)


@router.get("/sessions/{session_id}/transcript.docx")
def download_transcript_docx(session_id: str):
    """Build a .docx from the session's diarized transcript and stream it."""
    path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
    if not os.path.exists(path):
        raise HTTPException(404, detail="Session not found")
    with open(path, "r", encoding="utf-8") as f:
        session = json.load(f)

    doc = docx.Document()
    doc.add_heading("Session Transcript", level=1)
    meta = doc.add_paragraph()
    meta.add_run(f"Session ID: {session.get('id', session_id)}\n").italic = True
    if session.get("date"):
        meta.add_run(f"Date: {session['date']}\n").italic = True
    if session.get("speaker_count"):
        meta.add_run(f"Speakers: {session['speaker_count']}").italic = True

    turns = session.get("transcript_diarized") or []
    if turns:
        for turn in turns:
            p = doc.add_paragraph()
            p.add_run(f"{turn.get('speaker', 'Unknown')}: ").bold = True
            p.add_run(turn.get("text", ""))
    else:
        doc.add_paragraph(session.get("transcript_raw") or session.get("transcript") or "")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{session_id}_transcript.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/sessions/{session_id}/speakers/merge")
def merge_speaker(session_id: str, req: MergeSpeakerRequest):
    """Relabel every `from_name` turn as `into_name` and collapse adjacent same-speaker turns."""
    if req.from_name == req.into_name:
        raise HTTPException(400, detail="from_name and into_name are identical")

    path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
    if not os.path.exists(path):
        raise HTTPException(404, detail="Session not found")
    with open(path, "r", encoding="utf-8") as f:
        session = json.load(f)

    turns = session.get("transcript_diarized") or []
    found = any(t.get("speaker") == req.from_name for t in turns)
    if not found:
        raise HTTPException(404, detail=f"No turns found for speaker '{req.from_name}'")

    for t in turns:
        if t.get("speaker") == req.from_name:
            t["speaker"] = req.into_name

    turns = _collapse_adjacent_turns(turns)
    session = _save_turns(session, turns, path)
    return {
        "status": "ok",
        "transcript_diarized": session["transcript_diarized"],
        "speaker_count": session["speaker_count"],
    }


@router.post("/sessions/{session_id}/speakers/remove")
def remove_speaker(session_id: str, req: RemoveSpeakerRequest):
    """Delete every turn belonging to `name`, then collapse adjacent same-speaker turns."""
    path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
    if not os.path.exists(path):
        raise HTTPException(404, detail="Session not found")
    with open(path, "r", encoding="utf-8") as f:
        session = json.load(f)

    turns = session.get("transcript_diarized") or []
    remaining = [t for t in turns if t.get("speaker") != req.name]
    if len(remaining) == len(turns):
        raise HTTPException(404, detail=f"No turns found for speaker '{req.name}'")

    remaining = _collapse_adjacent_turns(remaining)
    session = _save_turns(session, remaining, path)
    return {
        "status": "ok",
        "transcript_diarized": session["transcript_diarized"],
        "speaker_count": session["speaker_count"],
    }


@router.post("/sessions/{session_id}/turns/{turn_index}")
def edit_turn(session_id: str, turn_index: int, req: EditTurnRequest):
    """Replace the text of a single turn by index. Keeps speaker and ordering intact."""
    new_text = (req.text or "").strip()
    if not new_text:
        raise HTTPException(400, detail="text must not be empty")

    path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
    if not os.path.exists(path):
        raise HTTPException(404, detail="Session not found")
    with open(path, "r", encoding="utf-8") as f:
        session = json.load(f)

    turns = session.get("transcript_diarized") or []
    if turn_index < 0 or turn_index >= len(turns):
        raise HTTPException(404, detail=f"Turn index {turn_index} out of range")

    turns[turn_index]["text"] = new_text
    session["transcript_diarized"] = turns
    with open(path, "w", encoding="utf-8") as f:
        json.dump(session, f, ensure_ascii=False, indent=2)

    return {"status": "ok", "transcript_diarized": turns}


@router.post("/sessions/{session_id}/speakers/rename")
def rename_speaker(session_id: str, req: RenameSpeakerRequest):
    """Rename a speaker label across every turn of the session."""
    new_name = (req.new_name or "").strip()
    if not new_name:
        raise HTTPException(400, detail="new_name must not be empty")
    if new_name == req.old_name:
        raise HTTPException(400, detail="new_name is identical to old_name")

    path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
    if not os.path.exists(path):
        raise HTTPException(404, detail="Session not found")

    with open(path, "r", encoding="utf-8") as f:
        session = json.load(f)

    turns = session.get("transcript_diarized") or []
    changed = 0
    for turn in turns:
        if turn.get("speaker") == req.old_name:
            turn["speaker"] = new_name
            changed += 1
    if changed == 0:
        raise HTTPException(404, detail=f"No turns found for speaker '{req.old_name}'")

    session = _save_turns(session, turns, path)
    return {
        "status": "ok",
        "renamed_turns": changed,
        "transcript_diarized": session["transcript_diarized"],
        "speaker_count": session["speaker_count"],
    }


@router.delete("/sessions/{session_id}")
def delete_session(session_id: str):
    path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
    if not os.path.exists(path):
        raise HTTPException(404, detail="Session not found")
    # Remove the persisted audio first so we don't leave orphans if the session delete fails.
    try:
        with open(path, "r", encoding="utf-8") as f:
            session = json.load(f)
        audio_filename = session.get("audio_filename")
        if audio_filename:
            audio_path = os.path.join(AUDIO_DIR, audio_filename)
            if os.path.exists(audio_path):
                os.remove(audio_path)
    except Exception:
        pass
    os.remove(path)
    return {"status": "deleted", "session_id": session_id}


@router.post("/sessions/{session_id}/transcript/override")
async def override_transcript(session_id: str, file: UploadFile = File(...)):
    """
    Override a session's transcript with an uploaded .docx or .pdf file.
    Re-runs diarization on the new content.
    """
    filename = file.filename.lower()

    path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
    if not os.path.exists(path):
        raise HTTPException(404, detail="Session not found")

    try:
        # Extract text from the uploaded file
        if filename.endswith(".docx"):
            source_type = "manual_docx"
            contents = file.file.read()
            doc = docx.Document(io.BytesIO(contents))
            new_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        elif filename.endswith(".pdf"):
            import fitz
            source_type = "manual_pdf"
            contents = file.file.read()
            pdf_doc = fitz.open(stream=contents, filetype="pdf")
            new_text = "\n".join(page.get_text() for page in pdf_doc)
        else:
            raise HTTPException(400, detail="Only .docx and .pdf files are supported")

        if not new_text.strip():
            raise HTTPException(400, detail="Could not extract text from the uploaded file")

        # Re-diarize the new transcript
        diarization = await run_in_threadpool(diarize_transcript, new_text)

        # Load and update session JSON
        with open(path, "r", encoding="utf-8") as f:
            session = json.load(f)

        session["source"] = source_type
        session["transcript_raw"] = new_text
        session["transcript_diarized"] = diarization["turns"]
        session["speaker_count"] = diarization["speaker_count"]
        # Keep legacy field for backwards compat
        session["transcript"] = new_text
        # Clear old analysis since transcript changed
        session["analysis"] = None

        # Write back
        with open(path, "w", encoding="utf-8") as f:
            json.dump(session, f, ensure_ascii=False, indent=2)

        return {
            "status": "success",
            "session_id": session_id,
            "source": source_type,
            "speaker_count": diarization["speaker_count"]
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, detail=f"Failed to process document: {str(e)}")
